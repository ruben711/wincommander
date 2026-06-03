"""Extraheert platte tekst uit alle cursusbestanden (pdf/docx/pptx/crdownload)
naar .source/*.txt zodat agents met exacte tekst werken i.p.v. afbeeldingen."""
import os, re, zipfile, sys

DL = r"C:\Users\ruben\Downloads"
OUT = r"C:\Users\ruben\WinCommander\.source"
os.makedirs(OUT, exist_ok=True)

FILES = [
    "Labo - Bestanden beheren met administrator credentials.docx",
    "Powershelloefeningoplossing.pdf", "Credentials.pdf",
    "Systemsmanagementsem2PowershellExamebOpgave.pdf",
    "ITIL.pdf", "ITIL Guiding Principles Test.pdf", "ITIL4_VIVES_Kortrijk.pptx",
    "Message Box in PowerShell gebruiken.pdf", "Gegevens opslaan met PS.pdf",
    "oplossingLoops.pdf", "Unconfirmed 244552.crdownload", "Unconfirmed 118161.crdownload",
    "OpdrachtPS2.pdf", "Variabelen in Powershell.pptx", "scripting2oplossingen.pdf",
    "Powershell WMIoplossingen.pdf", "Scripting2.pdf", "Powershell1 (1).pptx",
    "Variabelen in Powershell (2).pptx", "Powershell1.pdf", "CIM.pdf",
    "Loops in Powershell.pptx", "extra Oefeningen.pdf", "Oplossing LABO 1.pdf",
    "Opdracht.pdf", "PSDRIVE.pdf", "opdrachtOUTFILEenMessageBox.pdf",
    "Labo Loops.docx", "Powershell WMI.docx", "Labo 4 - Users csv import.docx",
    "Labo 1 PowerShell Scripting.docx", "Labo 3 -Basisoefening Powershell.docx",
]

def safe(name):
    return re.sub(r"[^A-Za-z0-9._ -]", "_", os.path.splitext(name)[0]).strip()

def pdf_text(path):
    import fitz
    doc = fitz.open(path)
    return "\n".join(f"--- pagina {i+1} ---\n{p.get_text()}" for i, p in enumerate(doc))

def docx_text(path):
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)

def pptx_text(path):
    ns = "{http://schemas.openxmlformats.org/drawingml/2006/main}t"
    import xml.etree.ElementTree as ET
    out = []
    with zipfile.ZipFile(path) as z:
        slides = sorted(
            [n for n in z.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", n)],
            key=lambda n: int(re.search(r"(\d+)", n).group(1)),
        )
        for i, s in enumerate(slides):
            root = ET.fromstring(z.read(s))
            texts = [e.text for e in root.iter(ns) if e.text]
            out.append(f"--- slide {i+1} ---\n" + "\n".join(texts))
    return "\n".join(out)

summary = []
for f in FILES:
    src = os.path.join(DL, f)
    if not os.path.exists(src):
        summary.append(f"MISSING  {f}"); continue
    ext = os.path.splitext(f)[1].lower()
    try:
        if ext == ".pdf":
            txt = pdf_text(src)
        elif ext == ".docx":
            txt = docx_text(src)
        elif ext == ".pptx":
            txt = pptx_text(src)
        else:  # crdownload (plain text scripts)
            txt = open(src, "r", encoding="utf-8", errors="replace").read()
        dst = os.path.join(OUT, safe(f) + ".txt")
        open(dst, "w", encoding="utf-8").write(txt)
        summary.append(f"OK {len(txt):>7}  {safe(f)}.txt")
    except Exception as e:
        summary.append(f"FAIL     {f}  -> {type(e).__name__}: {e}")

print("\n".join(summary))
